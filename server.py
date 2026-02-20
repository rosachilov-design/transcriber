"""
Transcriber Pro — Local Review Dashboard
Lightweight local server for reviewing cloud-transcribed results.
No GPU needed. Loads .json state files and pairs them with local audio.

Also supports S3 upload for sending files to RunPod.
"""

from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
import os
import sys
import io
import json
import re
import threading
import time
from pathlib import Path

import boto3
from botocore.config import Config
from docx import Document
from dotenv import load_dotenv

# Load .env credentials
load_dotenv()

# Fix UTF-8 on Windows
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Directories ───
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# ─── S3 Config ───
S3_BUCKET = "ez2d4o9xmt"
S3_ENDPOINT = "https://s3api-us-wa-1.runpod.io"
S3_REGION = "us-wa-1"

s3 = boto3.client(
    "s3",
    endpoint_url=S3_ENDPOINT,
    region_name=S3_REGION,
    aws_access_key_id=os.getenv("RUNPOD_ACCESS_KEY"),
    aws_secret_access_key=os.getenv("RUNPOD_SECRET_KEY"),
    config=Config(signature_version="s3v4"),
)

# ─── State ───
transcriptions = {}

def download_results_from_s3():
    """Check S3 for any finished results (.json, .md, .docx) and pull them to local uploads."""
    print("☁️ Checking for new results on RunPod Cloud...")
    try:
        response = s3.list_objects_v2(Bucket=S3_BUCKET)
        if 'Contents' not in response:
            return

        result_exts = {".json", ".md", ".docx"}
        for obj in response['Contents']:
            file_name = obj['Key']
            ext = Path(file_name).suffix.lower()
            
            if ext in result_exts:
                local_path = UPLOAD_DIR / file_name
                # Only download if it's new
                if not local_path.exists():
                    print(f"  📥 Downloading new result: {file_name}")
                    s3.download_file(S3_BUCKET, file_name, str(local_path))
                    
                    # If it's a JSON, load it into memory
                    if ext == ".json":
                        try:
                            with open(local_path, "r", encoding="utf-8") as f:
                                data = json.load(f)
                                task_id = data.get("filename")
                                if task_id:
                                    transcriptions[task_id] = data
                        except:
                            pass
    except Exception as e:
        print(f"⚠️ Cloud sync check failed: {e}")

def load_existing_tasks():
    """Load previously completed transcriptions from JSON files on disk."""
    print("📂 Scanning local uploads...")
    for json_file in UPLOAD_DIR.glob("*.json"):
        try:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                task_id = data.get("filename")
                if task_id:
                    transcriptions[task_id] = data
                    print(f"  ✅ Loaded: {task_id}")
        except Exception as e:
            print(f"  ❌ Failed to load {json_file.name}: {e}")

# Initial load from disk
load_existing_tasks()
# Immediate check for cloud results
download_results_from_s3()

# Start a background thread to check cloud every 30 seconds
def cloud_watchdog():
    while True:
        time.sleep(30)
        download_results_from_s3()

threading.Thread(target=cloud_watchdog, daemon=True).start()


# ─── Helpers ───

def format_timestamp(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02}:{minutes:02}:{secs:02}"
    return f"{minutes:02}:{secs:02}"


def clean_hallucinations(text: str) -> str:
    """Remove common Russian Whisper hallucinations."""
    hallucination_patterns = [
        r'\bРедактор субтитров\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bКорректор\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bСубтитры\s*:\s*[^\.]+',
        r'\bПеревод\s*:\s*[^\.]+',
        r'\bОзвучка\s*:\s*[^\.]+',
        r'\bРедактор субтитров\b',
        r'\bКорректор\b',
        r'\b(Все права защищены|Продолжение следует|Ставьте лайки|Подписывайтесь на канал)\b',
    ]
    cleaned = text
    for pattern in hallucination_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def generate_docx(task_id):
    """Generate a .docx file from the transcription segments."""
    if task_id not in transcriptions:
        return None
    task = transcriptions[task_id]
    file_path = UPLOAD_DIR / task["filename"]
    docx_file_path = file_path.with_suffix(".docx")

    doc = Document()
    doc.add_heading(f"Transcription: {task['filename']}", 0)
    for seg in task["result"]:
        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{seg['timestamp']}] {seg['speaker']}: ")
        ts_run.bold = True
        p.add_run(seg['text'])
    doc.save(docx_file_path)
    return docx_file_path.name


def regenerate_files(task_id):
    """Re-save .md, .docx, and .json after speaker edits."""
    task = transcriptions[task_id]
    file_path = UPLOAD_DIR / task["filename"]

    # MD
    md_file_path = file_path.with_suffix(".md")
    md_content = f"# Transcription: {task['filename']}\n\n"
    for seg in task["result"]:
        md_content += f"**[{seg['timestamp']}] {seg['speaker']}:** {seg['text']}\n\n"
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    # DOCX
    generate_docx(task_id)

    # JSON state
    state_file = file_path.with_suffix(".json")
    with open(state_file, "w", encoding="utf-8") as f:
        json.dump(task, f, indent=2, ensure_ascii=False)


# ─── S3 Upload (Background Thread) ───

def upload_to_s3(file_path: Path, task_id: str):
    """Upload audio file to RunPod S3 bucket in background."""
    try:
        transcriptions[task_id]["status"] = "uploading"
        transcriptions[task_id]["progress"] = 5

        file_size = file_path.stat().st_size
        uploaded = 0

        def progress_callback(bytes_transferred):
            nonlocal uploaded
            uploaded += bytes_transferred
            pct = min(int((uploaded / file_size) * 90), 90)
            transcriptions[task_id]["progress"] = pct

        s3.upload_file(
            str(file_path),
            S3_BUCKET,
            file_path.name,
            Callback=progress_callback,
        )

        transcriptions[task_id]["status"] = "uploaded"
        transcriptions[task_id]["progress"] = 100
        transcriptions[task_id]["s3_key"] = file_path.name
        print(f"☁️ Uploaded {file_path.name} to S3")

    except Exception as e:
        transcriptions[task_id]["status"] = "error"
        transcriptions[task_id]["error"] = f"S3 upload failed: {e}"
        print(f"❌ S3 upload failed: {e}")


# ═══════════════════════════════════════════
#  API ENDPOINTS
# ═══════════════════════════════════════════

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Save file locally and begin S3 upload in background."""
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    task_id = file.filename

    # Check if we already have a transcription for this file
    json_path = file_path.with_suffix(".json")
    if json_path.exists():
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                transcriptions[task_id] = data
                print(f"📎 Found existing transcription for {task_id}")
                return {"task_id": task_id}
        except:
            pass

    transcriptions[task_id] = {
        "filename": file.filename,
        "status": "uploading",
        "progress": 0,
        "result": [],
    }

    # Start S3 upload in background
    t = threading.Thread(target=upload_to_s3, args=(file_path, task_id), daemon=True)
    t.start()

    return {"task_id": task_id}


@app.get("/check/{filename}")
async def check_transcription(filename: str):
    """Check if a transcription JSON already exists for this audio file."""
    # Check in-memory first
    if filename in transcriptions and transcriptions[filename].get("status") == "completed":
        return transcriptions[filename]

    # Check on disk
    json_path = UPLOAD_DIR / Path(filename).with_suffix(".json")
    if json_path.exists():
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                transcriptions[filename] = data
                return data
        except:
            pass

    return {"status": "not_found"}


@app.get("/status/{task_id}")
async def get_status(task_id: str):
    """Return task status."""
    task = transcriptions.get(task_id)
    if not task:
        return {"status": "not_found"}
    return task


@app.get("/audio/{filename}")
async def get_audio(filename: str):
    """Serve audio file for the local player."""
    return FileResponse(UPLOAD_DIR / filename)


@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download .md or .docx result files."""
    path = UPLOAD_DIR / filename
    if path.exists():
        media_type = "text/markdown"
        if filename.endswith(".docx"):
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return FileResponse(path, media_type=media_type, filename=filename)
    return {"error": "File not found"}


@app.post("/save/{task_id}")
async def save_files(task_id: str):
    """Generate and save .md and .docx from the current transcription state."""
    if task_id not in transcriptions:
        return {"error": "Task not found"}

    task = transcriptions[task_id]
    if not task.get("result"):
        return {"error": "No transcription data to save"}

    regenerate_files(task_id)
    return {
        "status": "saved",
        "md_path": Path(task["filename"]).with_suffix(".md").name,
        "docx_path": Path(task["filename"]).with_suffix(".docx").name,
    }


class UpdateSpeakerRequest(BaseModel):
    task_id: str
    segment_index: int
    speaker_name: str


@app.post("/update_speaker")
async def update_speaker(req: UpdateSpeakerRequest):
    """Bulk rename a speaker across all segments."""
    if req.task_id in transcriptions:
        task = transcriptions[req.task_id]
        if 0 <= req.segment_index < len(task["result"]):
            old_name = task["result"][req.segment_index]["speaker"]
            new_name = req.speaker_name

            for seg in task["result"]:
                if seg["speaker"] == old_name:
                    seg["speaker"] = new_name

            regenerate_files(req.task_id)
            return {"status": "success"}

    return {"status": "error", "message": "Task or segment not found"}


@app.get("/list")
async def list_transcriptions():
    """List all available transcriptions (for a file picker)."""
    results = []
    for json_file in UPLOAD_DIR.glob("*.json"):
        try:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                results.append({
                    "filename": data.get("filename"),
                    "segments": len(data.get("result", [])),
                    "status": data.get("status", "unknown"),
                })
        except:
            pass
    return results


# ─── Static Files & Startup ───
app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    print(f"🚀 Transcriber Pro (Local Dashboard) starting on port {port}...")
    uvicorn.run(app, host="0.0.0.0", port=port)
