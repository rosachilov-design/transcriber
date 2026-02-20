"""
engine.py — GPU Transcription Engine
This file contains all the heavy GPU logic (Whisper + Diarization).
Runs ONLY on RunPod Pods with a GPU. Never runs locally.
"""

import whisper
import torch
import os
import json
import subprocess
import math
import re
import threading
from pathlib import Path

from docx import Document

# ─── Config ───
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
CACHE_DIR = Path("cache")
CACHE_DIR.mkdir(exist_ok=True)

device = "cuda" if torch.cuda.is_available() else "cpu"
model_lock = threading.Lock()

# ─── Load Models ───
print(f"🧠 Loading Whisper model on {device}...")
model = whisper.load_model("turbo", device=device)
print("✅ Whisper loaded.")

HF_TOKEN = "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"
print("🔊 Loading Diarization Pipeline...")
try:
    from pyannote.audio import Pipeline
    diarization_pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1",
        token=HF_TOKEN
    )
    if diarization_pipeline:
        diarization_pipeline.to(torch.device(device))
        print("✅ Diarization Pipeline loaded.")
    else:
        print("❌ Failed to load Diarization Pipeline.")
        diarization_pipeline = None
except Exception as e:
    print(f"❌ Error loading diarization pipeline: {e}")
    diarization_pipeline = None


# ─── Helpers ───

def format_timestamp(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02}:{minutes:02}:{secs:02}"
    return f"{minutes:02}:{secs:02}"


def get_duration(file_path):
    cmd = ["ffprobe", "-v", "error", "-show_entries", "format=duration",
           "-of", "default=noprint_wrappers=1:nokey=1", str(file_path)]
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
    return float(result.stdout.strip())


def clean_hallucinations(text: str) -> str:
    """Remove common Russian Whisper hallucinations."""
    hallucination_patterns = [
        r'\bРедактор субтитров\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bКорректор\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bСубтитры\s*:\s*[^\.]+',
        r'\bПеревод\s*:\s*[^\.]+',
        r'\bОзвучка\s*:\s*[^\.]+',
        r'\bРедактор субтитров\b',
        r'\bКорректор\b',
        r'\b(Все права защищены|Продолжение следует|Ставьте лайки|Подписывайтесь на канал)\b',
    ]
    cleaned = text
    for pattern in hallucination_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def get_speaker_for_word(timeline, word_start, word_end):
    """Assign a speaker to a word using diarization timeline."""
    if not timeline:
        return "Unknown"

    best_speaker = None
    best_overlap = 0

    for entry in timeline:
        overlap_start = max(word_start, entry["start"])
        overlap_end = min(word_end, entry["end"])
        if overlap_end > overlap_start:
            overlap = overlap_end - overlap_start
            if overlap > best_overlap:
                best_overlap = overlap
                best_speaker = entry["speaker"]

    if best_speaker:
        return best_speaker

    mid = (word_start + word_end) / 2
    nearest = min(timeline, key=lambda s: min(abs(s["start"] - mid), abs(s["end"] - mid)))
    return nearest["speaker"]


# ─── Diarization ───

def run_diarization(file_path: Path):
    if not diarization_pipeline:
        return []

    cache_file = CACHE_DIR / f"{file_path.stem}_diarize.json"
    if cache_file.exists():
        print(f"📂 Loading cached diarization for {file_path.name}...")
        with open(cache_file, 'r') as f:
            return json.load(f)

    print(f"🔊 Running diarization on {file_path.name}...")
    wav_path = CACHE_DIR / f"{file_path.stem}_diarize.wav"
    cmd = [
        "ffmpeg", "-y", "-i", str(file_path),
        "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", str(wav_path)
    ]
    subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    import soundfile as sf
    import numpy as np
    data, sample_rate = sf.read(str(wav_path), dtype='float32')
    if data.ndim == 1:
        data = data[np.newaxis, :]
    else:
        data = data.T
    waveform = torch.from_numpy(data)
    audio_input = {"waveform": waveform, "sample_rate": sample_rate}

    with model_lock:
        diarize_output = diarization_pipeline(audio_input, min_speakers=2)

    annotation = diarize_output.speaker_diarization if hasattr(diarize_output, 'speaker_diarization') else diarize_output

    timeline = []
    for turn, _, speaker in annotation.itertracks(yield_label=True):
        timeline.append({"start": turn.start, "end": turn.end, "speaker": speaker})

    with open(cache_file, 'w') as f:
        json.dump(timeline, f)

    if wav_path.exists():
        os.remove(wav_path)

    print(f"✅ Diarization complete: {len(timeline)} speaker turns.")
    return timeline


# ─── Main Transcription ───

def transcribe_file(file_path: Path, progress_callback=None):
    """
    Full transcription pipeline: diarization + chunked Whisper + speaker alignment.
    Returns a dict with 'result' (list of segments), 'filename', and 'status'.
    
    progress_callback(status, progress_pct) is called for live updates.
    """
    task = {
        "filename": file_path.name,
        "status": "diarizing",
        "progress": 5,
        "result": [],
    }

    def report(status, pct):
        task["status"] = status
        task["progress"] = pct
        if progress_callback:
            progress_callback(status, pct)

    try:
        # Phase 1: Diarization
        report("diarizing", 5)
        timeline = run_diarization(file_path)

        # Phase 2: Chunked transcription
        report("transcribing", 10)

        natural_chunks = []
        if timeline:
            current_chunk_turns = []
            chunk_start_time = timeline[0]["start"]

            for i, entry in enumerate(timeline):
                current_chunk_turns.append(entry)
                elapsed = entry["end"] - chunk_start_time
                is_last = (i == len(timeline) - 1)

                if elapsed >= 30 or is_last:
                    natural_chunks.append({
                        "start": chunk_start_time,
                        "end": entry["end"],
                        "turns": current_chunk_turns
                    })
                    if not is_last:
                        current_chunk_turns = []
                        chunk_start_time = timeline[i + 1]["start"]

        if not natural_chunks:
            duration = get_duration(file_path)
            natural_chunks = [{"start": i * 30, "end": min((i + 1) * 30, duration)}
                              for i in range(math.ceil(duration / 30))]

        speaker_map = {}
        speaker_counter = 0
        all_speaker_words = []

        for i, chunk in enumerate(natural_chunks):
            start_time = chunk["start"]
            end_time = chunk["end"]
            duration_s = end_time - start_time

            if duration_s <= 0:
                continue

            chunk_path = CACHE_DIR / f"{file_path.stem}_chunk_{i}.wav"

            cmd = [
                "ffmpeg", "-y", "-ss", str(start_time), "-t", str(duration_s),
                "-i", str(file_path), "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", str(chunk_path)
            ]
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

            if not chunk_path.exists():
                continue

            previous_context = ""
            if all_speaker_words:
                last_words = [sw["word"] for sw in all_speaker_words[-50:]]
                previous_context = " ".join(last_words)

            print(f"  📝 Chunk {i + 1}/{len(natural_chunks)} ({start_time:.1f}s - {end_time:.1f}s)")

            with model_lock:
                result = model.transcribe(
                    str(chunk_path),
                    language="ru",
                    verbose=False,
                    fp16=(device == "cuda"),
                    word_timestamps=True,
                    condition_on_previous_text=True,
                    initial_prompt=previous_context if previous_context else "Это аудиозапись беседы или интервью.",
                )

            for segment in result["segments"]:
                words = segment.get("words", [])
                if not words:
                    text = re.sub(r'\[.*?\]', '', segment["text"]).strip()
                    if text:
                        abs_start = segment["start"] + start_time
                        abs_end = segment["end"] + start_time
                        raw = get_speaker_for_word(timeline, abs_start, abs_end)
                        all_speaker_words.append({
                            "word": text, "start": abs_start,
                            "end": abs_end, "speaker_raw": raw
                        })
                    continue

                for w in words:
                    word_text = w.get("word", "").strip()
                    if not word_text:
                        continue
                    abs_start = w["start"] + start_time
                    abs_end = w["end"] + start_time
                    raw = get_speaker_for_word(timeline, abs_start, abs_end)
                    all_speaker_words.append({
                        "word": word_text, "start": abs_start,
                        "end": abs_end, "speaker_raw": raw
                    })

            if chunk_path.exists():
                os.remove(chunk_path)

            pct = 10 + int(((i + 1) / len(natural_chunks)) * 80)
            report("transcribing", pct)

        # Phase 3: Final grouping
        report("aligning", 95)

        final_segments = []
        if all_speaker_words:
            current_speaker_raw = all_speaker_words[0]["speaker_raw"]
            current_start = all_speaker_words[0]["start"]
            current_words = [all_speaker_words[0]["word"]]

            for sw in all_speaker_words[1:]:
                if sw["speaker_raw"] == current_speaker_raw:
                    current_words.append(sw["word"])
                else:
                    if current_speaker_raw not in speaker_map:
                        speaker_counter += 1
                        speaker_map[current_speaker_raw] = f"Speaker {speaker_counter}"
                    text = clean_hallucinations(" ".join(current_words))
                    text = re.sub(r'\s+', ' ', text)
                    if text:
                        final_segments.append({
                            "start": current_start,
                            "timestamp": format_timestamp(current_start),
                            "text": text,
                            "speaker": speaker_map[current_speaker_raw]
                        })
                    current_speaker_raw = sw["speaker_raw"]
                    current_start = sw["start"]
                    current_words = [sw["word"]]

            if current_words:
                if current_speaker_raw not in speaker_map:
                    speaker_counter += 1
                    speaker_map[current_speaker_raw] = f"Speaker {speaker_counter}"
                text = clean_hallucinations(" ".join(current_words))
                text = re.sub(r'\s+', ' ', text)
                if text:
                    final_segments.append({
                        "start": current_start,
                        "timestamp": format_timestamp(current_start),
                        "text": text,
                        "speaker": speaker_map[current_speaker_raw]
                    })

        # Merge adjacent same-speaker segments
        smoothed = []
        for seg in final_segments:
            if smoothed and seg["speaker"] == smoothed[-1]["speaker"]:
                smoothed[-1]["text"] += " " + seg["text"]
            else:
                smoothed.append(seg.copy())

        task["result"] = smoothed
        task["status"] = "completed"
        task["progress"] = 100

        print(f"✅ Transcription complete: {len(smoothed)} speaker turns.")
        return task

    except Exception as e:
        import traceback
        traceback.print_exc()
        task["status"] = "error"
        task["error"] = str(e)
        return task


def save_results(task, upload_dir=None):
    """Save .json, .md, and .docx files from a completed task."""
    if upload_dir is None:
        upload_dir = UPLOAD_DIR

    file_path = Path(upload_dir) / task["filename"]

    # Save .md
    md_path = file_path.with_suffix(".md")
    md_content = f"# Transcription: {task['filename']}\n\n"
    for seg in task["result"]:
        md_content += f"**[{seg['timestamp']}] {seg['speaker']}:** {seg['text']}\n\n"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    # Save .docx
    docx_path = file_path.with_suffix(".docx")
    doc = Document()
    doc.add_heading(f"Transcription: {task['filename']}", 0)
    for seg in task["result"]:
        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{seg['timestamp']}] {seg['speaker']}: ")
        ts_run.bold = True
        p.add_run(seg['text'])
    doc.save(docx_path)

    # Save .json (portable state)
    task["md_path"] = md_path.name
    task["docx_path"] = docx_path.name
    json_path = file_path.with_suffix(".json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(task, f, indent=2, ensure_ascii=False)

    print(f"💾 Saved: {md_path.name}, {docx_path.name}, {json_path.name}")
