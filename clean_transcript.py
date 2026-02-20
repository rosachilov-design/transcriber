import re
from docx import Document
from pathlib import Path

def clean_text(text):
    # 1. Remove Hallucinations
    hallucinations = [
        r"Субтитры делал DimaTorzok",
        r"Субтитры создавал DimaTorzok",
        r"anomaly TikTok\.",
        r"Пока\. Пока\. Пока\. Почему\?",
        r"Когда я кушаю лампе оперный, когда яagram лайк\. Я считает ,riet о том, что это происходит, такое тв\.",
        r"на приходите fixed 두 세間тр енток условия\.",
        r"Квень памяти, знаете, у машины, смей, в помощь сちゃущее failёвкиатринFL в режиме\.",
        r"Уг votes please anybody to\.\.\.",
        r"Bible ночам",
        r"сделать siamo Sabah",
        r"за 누ушеводного надто",
        r"Теперь у нас есть аспектная задача! Но сейчас вот как -то mobile проект\.",
        r"anomaly TikTok",
        r"igram лайк",
        r"siamo Sabah"
    ]
    for h in hallucinations:
        text = re.sub(h, "", text, flags=re.IGNORECASE)

    # 2. Fix Medical Terms & Drugs
    replacements = {
        r"гипертези": "гипертензи",
        r"Форстигу": "Форсигу",
        r"Симовик": "Семавик",
        r"Терезетта": "Тирзепатид",
        r"Идарби Кло": "Эдарби Кло",
        r"МХВ-10": "МКБ-10",
        r"третинопатия": "ретинопатия",
        r"поглифозин": "дапаглифлозин",
        r"диалтация": "дилатация",
        r"вал сортан": "валсартан",
        r"телми сортан": "телмисартан",
        r"азил сортан": "азилсартан",
        r"КПТН": "Капотен"
    }
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text)

    # Clean up multiple spaces left by removals
    text = re.sub(r' +', ' ', text)
    text = re.sub(r' ,', ',', text)
    text = re.sub(r' \.', '.', text)
    return text.strip()

def create_docx(input_file, output_file):
    doc = Document()
    doc.add_heading('Транскрипт интервью (Очищенная версия)', 0)
    
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Rename speakers for clarity
    content = content.replace("Speaker 1", "Модератор (Анастасия)")
    content = content.replace("Speaker 2", "Наталья (Интервьюер)")
    content = content.replace("Speaker 3", "Анна Юрьевна (Эксперт)")
    content = content.replace("Speaker 4", "Никита (Интервьюер)")

    # Splitting by segments (assuming segments start with **[XX:XX]**)
    segments = re.split(r'(\*\*\[\d{2}:\d{2}\].*?\*\*)', content)
    
    current_speaker_line = ""
    for part in segments:
        if re.match(r'\*\*\[\d{2}:\d{2}\].*?\*\*', part):
            current_speaker_line = part
        else:
            if current_speaker_line:
                text = clean_text(part)
                if text:
                    p = doc.add_paragraph()
                    # Add timestamp and speaker in bold
                    run = p.add_run(current_speaker_line + " ")
                    run.bold = True
                    p.add_run(text)
                current_speaker_line = ""

    doc.save(output_file)
    print(f"File saved to {output_file}")

if __name__ == "__main__":
    input_path = Path("uploads/Interview1.md")
    output_path = Path("uploads/Interview1_Cleaned.docx")
    create_docx(input_path, output_path)
